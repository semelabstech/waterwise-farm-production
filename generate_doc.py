from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    hs.font.name = 'Calibri'

# ── Helper functions ────────────────────────────────────
def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    run = p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Système de Précision Irrigation\npar Vision Satellite et IoT")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Plan d'Implémentation Technique")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x4C, 0xAF, 0x50)

doc.add_paragraph()

theme = doc.add_paragraph()
theme.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = theme.add_run("Thème : Agriculture & Stress Hydrique")
run.font.size = Pt(14)
run.italic = True

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("Avril 2026")
run.font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  TABLE OF CONTENTS placeholder
# ══════════════════════════════════════════════════════════
add_heading("Table des Matières", 1)
toc_items = [
    "1. Introduction et Contexte",
    "2. Problématique Scientifique",
    "3. Objectif Général",
    "4. Architecture du Projet",
    "5. Analyse des Fonctionnalités (F1–F6)",
    "6. Pipeline de Données (F1 + F2)",
    "7. Modèles Deep Learning (F3 + F4)",
    "8. Moteur de Fusion (F6)",
    "9. Dashboard Streamlit",
    "10. Stack Technique",
    "11. Plan de Vérification",
    "12. Conclusion",
]
for item in toc_items:
    add_para(item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ══════════════════════════════════════════════════════════
add_heading("1. Introduction et Contexte", 1)

add_para(
    "Le Maroc traverse une période de sécheresse historique qui menace directement "
    "le secteur agricole, principal consommateur de ressources en eau du pays. "
    "L'intelligence artificielle offre une opportunité unique d'optimiser chaque "
    "goutte d'eau grâce à l'analyse de données satellites, météorologiques et "
    "de capteurs de terrain."
)

add_heading("Problème Principal", 2)
add_para("L'irrigation agricole traditionnelle souffre de plusieurs limitations :")
problems = [
    "Uniforme — même quantité d'eau partout, sans distinction entre les parcelles",
    "Basée sur l'intuition — décisions sans données objectives",
    "Non adaptée aux conditions climatiques réelles en temps réel",
    "Incapable d'anticiper les variations météorologiques",
]
for p in problems:
    add_bullet(p)

add_heading("Conséquences", 2)
consequences = [
    "Gaspillage d'eau considérable",
    "Baisse de rendement des cultures",
    "Stress hydrique non détecté à temps",
    "Mauvaise planification de l'irrigation",
]
for c in consequences:
    add_bullet(c)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  2. PROBLÉMATIQUE
# ══════════════════════════════════════════════════════════
add_heading("2. Problématique Scientifique", 1)

add_para(
    "Comment exploiter des images satellites gratuites (Sentinel-2), combinées à "
    "des données météo et des capteurs low-cost d'humidité du sol, pour :"
)
questions = [
    "Détecter le stress hydrique à l'échelle d'une parcelle agricole",
    "Anticiper le besoin en eau à 24–48 heures",
    "Recommander une décision d'irrigation précise (où / quand / combien)",
]
for i, q in enumerate(questions, 1):
    add_bullet(f"{i}. {q}")

# ══════════════════════════════════════════════════════════
#  3. OBJECTIF
# ══════════════════════════════════════════════════════════
add_heading("3. Objectif Général", 1)

add_para(
    "Développer un système intelligent utilisant des images satellites (Sentinel-2), "
    "des capteurs IoT, des modèles de Deep Learning et des prévisions météo "
    "afin d'optimiser l'irrigation et réduire la consommation d'eau de 30%.",
    bold=True
)

components = [
    "Images satellites Sentinel-2 (Copernicus, gratuit)",
    "Capteurs IoT d'humidité du sol (ESP32 / Arduino)",
    "Modèles de Deep Learning (U-Net, Vision Transformer)",
    "Modèles de séries temporelles (Informer / Mamba / LSTM)",
    "Prévisions météorologiques (NASA POWER, Open-Meteo)",
]
add_para("Composants clés du système :")
for c in components:
    add_bullet(c)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  4. ARCHITECTURE DU PROJET
# ══════════════════════════════════════════════════════════
add_heading("4. Architecture du Projet", 1)

add_heading("Structure des Fichiers", 2)

structure = """
Theme Agriculture & Stress Hydrique/
├── config/
│   └── settings.py
├── data/
│   ├── satellite/
│   ├── weather/
│   ├── iot/
│   ├── processed/
│   └── synthetic/
├── models/
│   ├── unet/
│   │   ├── model.py
│   │   ├── dataset.py
│   │   └── train.py
│   ├── timeseries/
│   │   ├── model.py
│   │   ├── dataset.py
│   │   └── train.py
│   └── checkpoints/
├── pipeline/
│   ├── satellite.py
│   ├── weather.py
│   ├── iot.py
│   ├── indices.py
│   └── fusion.py
├── dashboard/
│   ├── app.py
│   ├── pages/
│   │   ├── map_view.py
│   │   ├── stress_analysis.py
│   │   ├── predictions.py
│   │   └── recommendations.py
│   └── components/
│       ├── sidebar.py
│       ├── charts.py
│       └── metrics.py
├── utils/
│   ├── geo.py
│   ├── visualization.py
│   └── evaluation.py
├── demo/
│   └── generate_synthetic.py
├── requirements.txt
├── README.md
└── run.py
""".strip()

p = doc.add_paragraph()
run = p.add_run(structure)
run.font.name = 'Consolas'
run.font.size = Pt(8)

add_heading("Architecture Système", 2)

add_para(
    "Le système suit une architecture modulaire en 4 couches :"
)
layers = [
    "Couche Données : Acquisition satellite (Copernicus API), météo (NASA POWER), IoT (MQTT/CSV)",
    "Couche Traitement : Calcul d'indices (NDVI, NDMI), préparation de patches 256×256, normalisation",
    "Couche Intelligence : U-Net/ViT pour segmentation du stress + Informer/LSTM pour prévision ET0",
    "Couche Décision : Fusion des 3 sources → recommandations d'irrigation ciblées",
]
for l in layers:
    add_bullet(l)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  5. FONCTIONNALITÉS
# ══════════════════════════════════════════════════════════
add_heading("5. Analyse des Fonctionnalités (F1–F6)", 1)

add_table(
    ["#", "Fonctionnalité", "Description", "Technologie"],
    [
        ["F1", "Sélection zone & période",
         "Définir zone agricole (bbox/polygone), période, filtre nuages → images Sentinel-2 L2A",
         "Copernicus API, geopandas"],
        ["F2", "Préparation données satellite",
         "Extraction bandes B04/B08/B11, calcul NDVI/NDMI, masquage nuages, patches 256×256",
         "rasterio, numpy"],
        ["F3", "Détection stress hydrique",
         "Segmentation U-Net/ViT : zones normales vs stress (binaire ou multi-niveaux)",
         "PyTorch, segmentation_models_pytorch"],
        ["F4", "Prévision besoin hydrique 48h",
         "Informer/LSTM sur séries temporelles météo → prédire ET0 ou risque futur",
         "PyTorch, scikit-learn"],
        ["F5", "Intégration capteurs IoT",
         "Mesurer humidité réelle du sol, validation terrain, correction satellite",
         "ESP32, MQTT, SQLite"],
        ["F6", "Fusion intelligente",
         "Combiner carte stress + prévision météo + humidité capteurs → décision irrigation",
         "Python, algorithme pondéré"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  6. PIPELINE DE DONNÉES
# ══════════════════════════════════════════════════════════
add_heading("6. Pipeline de Données (F1 + F2)", 1)

add_heading("6.1 Acquisition Satellite (F1)", 2)
add_para("Module : pipeline/satellite.py", italic=True)
add_para("Classe SentinelDownloader :")
steps_f1 = [
    "Authentification via Copernicus Data Space API (OIDC token)",
    "Recherche de produits Sentinel-2 L2A par bounding box, plage de dates, couverture nuageuse (<20%)",
    "Téléchargement automatique des fichiers .jp2 (bandes B04-Red, B08-NIR, B11-SWIR)",
    "Gestion du cache local pour éviter les re-téléchargements",
]
for s in steps_f1:
    add_bullet(s)

add_heading("6.2 Prétraitement Satellite (F2)", 2)
add_para("Classe SatellitePreprocessor :")
steps_f2 = [
    "Lecture des bandes avec rasterio (format .jp2)",
    "Application du masque de nuages SCL (Scene Classification Layer)",
    "Calcul NDVI = (B08 - B04) / (B08 + B04)",
    "Calcul NDMI = (B08 - B11) / (B08 + B11)",
    "Découpage en patches 256×256 pixels avec chevauchement optionnel",
    "Normalisation des valeurs [0, 1]",
    "Sauvegarde en format NumPy (.npy) structuré",
]
for s in steps_f2:
    add_bullet(s)

add_heading("6.3 Indices de Végétation", 2)
add_para("Module : pipeline/indices.py", italic=True)

add_table(
    ["Indice", "Formule", "Seuil Stress", "Interprétation"],
    [
        ["NDVI", "(NIR - Red) / (NIR + Red)", "< 0.3", "Végétation en stress / sol nu"],
        ["NDMI", "(NIR - SWIR) / (NIR + SWIR)", "< 0.0", "Déficit hydrique de la végétation"],
    ]
)

add_para("Classification multi-niveaux du stress :")
add_table(
    ["Niveau", "Code", "NDVI", "Action"],
    [
        ["Normal", "0", "> 0.5", "Pas d'irrigation nécessaire"],
        ["Stress léger", "1", "0.3 – 0.5", "Surveillance renforcée"],
        ["Stress modéré", "2", "0.2 – 0.3", "Irrigation recommandée"],
        ["Stress sévère", "3", "< 0.2", "Irrigation urgente"],
    ]
)

add_heading("6.4 Données Météorologiques", 2)
add_para("Module : pipeline/weather.py", italic=True)
add_para("Sources de données :")
weather_sources = [
    "NASA POWER API — données historiques et prévisionnelles (température, humidité, vent, précipitations, radiation solaire)",
    "Open-Meteo — alternative gratuite sans clé API",
    "Meteostat — données météo historiques régionales",
]
for w in weather_sources:
    add_bullet(w)

add_para(
    "Calcul de l'évapotranspiration de référence (ET0) selon la méthode FAO Penman-Monteith, "
    "standard international pour estimer les besoins en eau des cultures."
)

add_heading("6.5 Données IoT (F5)", 2)
add_para("Module : pipeline/iot.py", italic=True)
iot_items = [
    "Simulateur IoT : génère des données réalistes d'humidité du sol corrélées à la météo",
    "Format : timestamp, latitude, longitude, humidité (%)",
    "Stockage en CSV ou SQLite pour prototypage",
    "Support MQTT pour intégration de capteurs réels (ESP32/Arduino)",
    "Interpolation des données manquantes et nettoyage automatique",
]
for item in iot_items:
    add_bullet(item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  7. MODÈLES DEEP LEARNING
# ══════════════════════════════════════════════════════════
add_heading("7. Modèles Deep Learning (F3 + F4)", 1)

add_heading("7.1 U-Net — Détection du Stress Hydrique (F3)", 2)
add_para("Module : models/unet/model.py", italic=True)

add_para("Architecture principale :", bold=True)
unet_details = [
    "Encodeur : ResNet34 pré-entraîné sur ImageNet",
    "Décodeur : U-Net avec skip connections",
    "Input : patches 256×256 à 3 canaux (NDVI, NDMI, B04-Red)",
    "Output : masque de segmentation (binaire ou multi-classes)",
    "Bibliothèque : segmentation_models_pytorch (smp)",
]
for d in unet_details:
    add_bullet(d)

add_para("Alternative avancée — Vision Transformer (SegFormer) :", bold=True)
vit_details = [
    "Modèle SegFormer de la bibliothèque timm",
    "Capture des relations spatiales à longue distance",
    "Meilleure compréhension globale de la parcelle",
    "Innovation 2026 — état de l'art en segmentation",
]
for d in vit_details:
    add_bullet(d)

add_para("Entraînement :", bold=True)
train_unet = [
    "Fonction de perte : Dice Loss + Cross-Entropy (combinée)",
    "Optimiseur : AdamW, learning rate = 1e-4",
    "Augmentation : rotations, flips, ajustement de luminosité",
    "Métriques : IoU (Intersection over Union), F1-Score, Dice Coefficient",
    "Early stopping sur la validation IoU",
    "Mixed precision (FP16) pour accélérer l'entraînement",
]
for t in train_unet:
    add_bullet(t)

add_heading("7.2 Informer / LSTM — Prévision à 48h (F4)", 2)
add_para("Module : models/timeseries/model.py", italic=True)

add_para("Modèle Baseline — LSTM/GRU :", bold=True)
lstm_details = [
    "Input : séquence de 14 jours (température, humidité, vent, précipitations, radiation)",
    "Architecture : 2 couches LSTM bidirectionnelles + couche dense",
    "Output : prédiction ET0 à 24h et 48h",
    "Simple et robuste pour le prototypage",
]
for d in lstm_details:
    add_bullet(d)

add_para("Modèle Avancé — Informer :", bold=True)
informer_details = [
    "ProbSparse Self-Attention : réduit la complexité de O(L²) à O(L log L)",
    "Distilling layers : compression progressive de l'information",
    "Multi-head attention avec mécanisme d'attention probabiliste",
    "Supérieur au Transformer classique pour les longues séries temporelles",
    "Adapté aux prévisions météo avec saisonnalité",
]
for d in informer_details:
    add_bullet(d)

add_para("Entraînement :", bold=True)
train_ts = [
    "Fonction de perte : MSE + MAE (combinée)",
    "Optimiseur : AdamW avec scheduling cosine",
    "Normalisation : StandardScaler sur chaque feature",
    "Split temporel : 70% train, 15% validation, 15% test",
    "Métriques : RMSE, MAE, R²",
]
for t in train_ts:
    add_bullet(t)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  8. MOTEUR DE FUSION
# ══════════════════════════════════════════════════════════
add_heading("8. Moteur de Fusion Intelligent (F6)", 1)
add_para("Module : pipeline/fusion.py", italic=True)

add_para(
    "Le moteur de décision combine les trois sources de données pour produire "
    "des recommandations d'irrigation personnalisées par zone :"
)

add_para("Sources fusionnées :", bold=True)
fusion_sources = [
    "Carte de stress actuelle — issue du modèle U-Net / ViT (état présent de la parcelle)",
    "Prévision météo 48h — issue du modèle Informer / LSTM (besoin futur en eau)",
    "Humidité du sol — issue des capteurs IoT (vérité terrain)",
]
for i, s in enumerate(fusion_sources, 1):
    add_bullet(f"{i}. {s}")

add_para("Algorithme de fusion pondérée :", bold=True)
add_para(
    "Score_irrigation = w1 × stress_satellite + w2 × deficit_predit + w3 × (1 - humidite_sol)",
    indent=True
)
add_para(
    "Avec w1=0.4, w2=0.35, w3=0.25 (ajustables). "
    "Le score est converti en recommandation concrète :"
)

add_table(
    ["Score", "Priorité", "Recommandation", "Volume (mm)"],
    [
        ["< 0.3", "Basse", "Pas d'irrigation", "0"],
        ["0.3 – 0.5", "Moyenne", "Irrigation légère", "5 – 10"],
        ["0.5 – 0.7", "Haute", "Irrigation standard", "15 – 25"],
        ["> 0.7", "Urgente", "Irrigation intensive", "30 – 50"],
    ]
)

add_para("Sortie par zone :", bold=True)
outputs = [
    "Identifiant de zone",
    "Niveau de stress actuel",
    "Volume d'eau recommandé (mm)",
    "Priorité (basse / moyenne / haute / urgente)",
    "Fenêtre temporelle optimale d'irrigation",
    "Économie d'eau estimée vs irrigation uniforme",
]
for o in outputs:
    add_bullet(o)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  9. DASHBOARD
# ══════════════════════════════════════════════════════════
add_heading("9. Dashboard Streamlit", 1)

add_para(
    "Interface web interactive développée avec Streamlit, offrant une visualisation "
    "complète du système avec thème sombre professionnel."
)

add_heading("9.1 Pages du Dashboard", 2)

pages = [
    ("Carte Interactive (map_view.py)",
     "Carte Folium du Maroc avec sélection de parcelles, overlay NDVI/NDMI/stress en heatmap, zones colorées par niveau de stress"),
    ("Analyse du Stress (stress_analysis.py)",
     "Résultats de segmentation U-Net, distribution du stress (diagrammes circulaires, histogrammes), évolution temporelle du NDVI"),
    ("Prévisions 48h (predictions.py)",
     "Graphiques de prévision météo, courbe de prédiction ET0, timeline de risque avec intervalles de confiance"),
    ("Recommandations (recommendations.py)",
     "Tableau de planification d'irrigation, volume par zone, classement par priorité, économie d'eau estimée (% réduction), export CSV/PDF"),
]

for title, desc in pages:
    add_para(title, bold=True)
    add_para(desc, indent=True)

add_heading("9.2 Composants Réutilisables", 2)
components = [
    "Sidebar : sélecteur de date, sélecteur de zone, choix du modèle (U-Net/ViT)",
    "Charts : graphiques stylisés Plotly/Matplotlib avec palette cohérente",
    "Metrics : cartes KPI (eau économisée, stress détecté, précision du modèle)",
]
for c in components:
    add_bullet(c)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  10. STACK TECHNIQUE
# ══════════════════════════════════════════════════════════
add_heading("10. Stack Technique Complète", 1)

add_table(
    ["Catégorie", "Technologies", "Rôle"],
    [
        ["Données Satellite", "Copernicus API, Sentinel Hub SDK, rasterio, geopandas", "Acquisition et lecture images Sentinel-2"],
        ["Indices de Végétation", "numpy", "Calcul NDVI, NDMI"],
        ["Données Météo", "NASA POWER API, Open-Meteo, Meteostat, pandas", "Séries temporelles météorologiques"],
        ["Deep Learning Vision", "PyTorch, segmentation_models_pytorch, timm", "U-Net, SegFormer (Vision Transformer)"],
        ["Séries Temporelles", "PyTorch (LSTM, GRU, Informer), scikit-learn", "Prévision ET0 à 24-48h"],
        ["IoT", "ESP32 / Arduino, MQTT, SQLite / CSV", "Capteurs d'humidité du sol"],
        ["Évaluation", "torchmetrics, scikit-learn", "IoU, F1, Dice, MAE, RMSE, R²"],
        ["Visualisation", "Streamlit, matplotlib, seaborn, folium, plotly", "Dashboard interactif et cartographie"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  11. VÉRIFICATION
# ══════════════════════════════════════════════════════════
add_heading("11. Plan de Vérification", 1)

add_heading("11.1 Mode Démo", 2)
add_para(
    "Le système inclut un générateur de données synthétiques (demo/generate_synthetic.py) "
    "permettant de tester l'ensemble du pipeline sans accès aux API externes."
)
demo_steps = [
    "python run.py --mode demo → génère données synthétiques et lance toutes les étapes",
    "python -m models.unet.train --epochs 5 --demo → vérifie le pipeline d'entraînement U-Net",
    "python -m models.timeseries.train --epochs 5 --demo → vérifie le pipeline séries temporelles",
    "streamlit run dashboard/app.py → lance le dashboard interactif",
]
for s in demo_steps:
    add_bullet(s)

add_heading("11.2 Métriques de Validation", 2)
add_table(
    ["Modèle", "Métrique", "Objectif Minimum"],
    [
        ["U-Net (Segmentation)", "IoU (Intersection over Union)", "> 0.65"],
        ["U-Net (Segmentation)", "F1-Score", "> 0.70"],
        ["Informer (Prévision)", "RMSE (ET0)", "< 0.5 mm/jour"],
        ["Informer (Prévision)", "R²", "> 0.85"],
        ["Système Global", "Réduction eau estimée", "≥ 30%"],
    ]
)

add_heading("11.3 Vérification Manuelle", 2)
manual = [
    "Inspection visuelle des cartes NDVI/stress sur le dashboard",
    "Vérification que les zones stressées reçoivent plus d'eau dans les recommandations",
    "Cohérence entre prévisions météo et recommandations d'irrigation",
    "Test de l'export CSV/PDF des recommandations",
]
for m in manual:
    add_bullet(m)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  12. CONCLUSION
# ══════════════════════════════════════════════════════════
add_heading("12. Conclusion", 1)

add_para(
    "Ce projet de Système Intelligent d'Irrigation de Précision propose une solution "
    "innovante pour optimiser l'utilisation de l'eau agricole au Maroc face à la sécheresse. "
    "En combinant images satellites Sentinel-2, données météo, capteurs IoT et modèles "
    "d'Intelligence Artificielle (U-Net, Informer/Mamba), le système permet de :"
)

conclusions = [
    "Détecter le stress hydrique en temps réel à l'échelle de la parcelle",
    "Anticiper les besoins en eau à 24–48 heures avec précision",
    "Recommander une irrigation ciblée et personnalisée par zone",
    "Réduire la consommation d'eau d'environ 30% par rapport à l'irrigation uniforme",
]
for c in conclusions:
    add_bullet(c)

add_para("")
add_para(
    "Cette approche intégrée offre une réponse technologique durable et adaptée "
    "aux enjeux climatiques actuels, contribuant à la sécurité alimentaire et à la "
    "préservation des ressources hydriques du Maroc.",
    bold=True
)

# ── Save ────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), "Plan_Implementation_Irrigation_Precision.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
