import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_premium_doc():
    md_path = r"C:\Users\G1325\.gemini\antigravity\brain\3de1467b-8219-4da4-a479-203a73a73339\project_documentation.md"
    out_path = r"c:\Users\G1325\Downloads\Theme Agriculture & Stress Hydrique\DOCUMENTATION_FINALE_IA.docx"
    
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    doc = Document()
    
    # --- SETUP MARGINS ---
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- SETUP STYLES ---
    # Normal Text
    st_normal = doc.styles['Normal']
    st_normal.font.name = 'Segoe UI'
    st_normal.font.size = Pt(11)
    st_normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Dark Gray
    st_normal.paragraph_format.space_after = Pt(10)
    st_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    st_normal.paragraph_format.line_spacing = 1.15

    # Bullet lists
    # Note: python-docx has List Bullet
    
    # Headings
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Segoe UI Semibold'
    h1.font.size = Pt(22)
    h1.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20) # Deep Green
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Segoe UI Semibold'
    h2.font.size = Pt(16)
    h2.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32) # Green
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    h3 = doc.styles['Heading 3']
    h3.font.name = 'Segoe UI'
    h3.font.size = Pt(13)
    h3.font.color.rgb = RGBColor(0x38, 0x8E, 0x3C)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    # --- COVER PAGE ---
    for _ in range(4):
        doc.add_paragraph()
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("Système d'Irrigation de Précision par IA")
    run_title.font.size = Pt(28)
    run_title.font.name = 'Segoe UI'
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    
    doc.add_paragraph()
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Documentation Technique & Pipeline d'Intelligence Artificielle")
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    for _ in range(2):
        doc.add_paragraph()
        
    p_theme = doc.add_paragraph()
    p_theme.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_theme = p_theme.add_run("Thème : Agriculture & Stress Hydrique — Maroc")
    run_theme.font.size = Pt(14)
    run_theme.bold = True
    
    for _ in range(8):
        doc.add_paragraph()
        
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run("Version Finale - 2026")
    run_date.font.size = Pt(12)
    run_date.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    
    doc.add_page_break()

    # --- PARSER VARIABLES ---
    lines = md_text.split('\n')
    in_code_block = False
    table_lines = []
    
    def process_bold_text(paragraph, text):
        # Splits by **
        parts = text.split('**')
        for i, part in enumerate(parts):
            if not part: continue
            run = paragraph.add_run(part)
            if i % 2 != 0: # Odd indices are inside **...**
                run.bold = True

    def flush_table():
        if not table_lines:
            return
            
        rows = []
        for line in table_lines:
            if line.strip().startswith('|'):
                row = [cell.strip() for cell in line.split('|')[1:-1]]
                # check if separator
                if all(c.replace('-', '').strip() == '' for c in row):
                    continue
                rows.append(row)
        
        if rows:
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Medium Shading 1 Accent 3' # A nice modern professional table style built into Word
            table.autofit = True
            
            for i, row in enumerate(rows):
                for j, cell_text in enumerate(row):
                    if j < len(table.columns):
                        cell = table.cell(i, j)
                        # Clean cell text
                        clean_text = cell_text.replace('**', '')
                        cell.text = clean_text
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for r in p.runs:
                                r.font.name = 'Segoe UI'
                                r.font.size = Pt(10)
                                if i == 0:
                                    r.bold = True
        table_lines.clear()
        doc.add_paragraph()

    # --- PARSING LOOP ---
    for line in lines:
        stripped = line.strip()
        
        # Skip Mermaid & Code
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            continue
            
        # Tables
        if stripped.startswith('|'):
            table_lines.append(stripped)
            continue
        else:
            flush_table()
            
        # Empty lines
        if not stripped:
            continue # We let space_before/after handle spacing
            
        # Skip Markdown TOC
        if "Table des Matières" in stripped and stripped.startswith("## "):
            continue
        if stripped.startswith("1. [") or stripped.startswith("2. ["):
            continue # Skip markdown toc lines
            
        if stripped == '---':
            doc.add_paragraph()
            continue

        # Headings
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:].replace('**', ''), 4)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:].replace('**', ''), 3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:].replace('**', ''), 2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:].replace('**', ''), 1)
            
        # Images
        elif stripped.startswith('![') and '](' in stripped:
            try:
                img_path = stripped.split('](')[1].split(')')[0]
                caption_text = stripped.split('![')[1].split(']')[0]
                
                doc.add_paragraph() # Spacing
                
                if os.path.exists(img_path):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Make it big and nice
                    p_img.add_run().add_picture(img_path, width=Inches(6.5))
                    
                    # Caption
                    if caption_text:
                        p_cap = doc.add_paragraph()
                        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_cap = p_cap.add_run(caption_text)
                        run_cap.font.size = Pt(9)
                        run_cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                        run_cap.italic = True
                
                doc.add_paragraph() # Spacing
            except Exception as e:
                pass
                
        # Bullet points
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            process_bold_text(p, stripped[2:])
            
        # Quotes / Alerts
        elif stripped.startswith('> '):
            clean = stripped[2:]
            if clean.startswith('[!'): # Github alerts
                clean = clean.split(']')[-1].strip()
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            # Add a subtle background color via borders or just make it italic
            run = p.add_run("📝 " + clean)
            run.italic = True
            run.font.color.rgb = RGBColor(0x00, 0x52, 0x9B) # Deep blue
            
        # Normal text
        else:
            # Check for math
            if stripped.startswith('$$') and stripped.endswith('$$'):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped.replace('$$', '').strip())
                run.font.name = 'Cambria Math'
                run.font.size = Pt(12)
            else:
                p = doc.add_paragraph()
                process_bold_text(p, stripped)

    doc.save(out_path)
    print("SUCCESS")

if __name__ == '__main__':
    create_premium_doc()
