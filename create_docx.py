import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_formatted_doc():
    md_path = r"C:\Users\G1325\.gemini\antigravity\brain\3de1467b-8219-4da4-a479-203a73a73339\project_documentation.md"
    
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    doc = Document()
    
    # Setup styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    for level in range(1, 5):
        try:
            hs = doc.styles[f'Heading {level}']
            hs.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
            hs.font.name = 'Calibri'
        except:
            pass

    # Helper for bold text
    def add_formatted_para(text, style_name='Normal'):
        p = doc.add_paragraph(style=style_name)
        
        # Simple bold regex
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        return p

    lines = md_text.split('\n')
    
    in_code_block = False
    table_lines = []
    
    def render_table():
        if not table_lines:
            return
        
        # parse table lines
        rows = []
        for line in table_lines:
            if line.strip().startswith('|'):
                row = [cell.strip() for cell in line.split('|')[1:-1]]
                if all(c.replace('-', '').strip() == '' for c in row):
                    continue # skip separator
                rows.append(row)
        
        if rows:
            # Create docx table
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            for i, row in enumerate(rows):
                for j, cell_text in enumerate(row):
                    if j < len(table.columns):
                        cell = table.cell(i, j)
                        # clean cell text
                        clean_text = cell_text.replace('**', '')
                        cell.text = clean_text
                        if i == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
        table_lines.clear()
        doc.add_paragraph()

    for line in lines:
        stripped = line.strip()
        
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            continue # Skip code/mermaid rendering in the doc for a cleaner look
            
        if stripped.startswith('|'):
            table_lines.append(stripped)
            continue
        else:
            render_table()
            
        if not stripped:
            doc.add_paragraph()
            continue
            
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:], 4)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], 3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], 2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], 1)
        elif stripped.startswith('![') and '](' in stripped:
            # Image extraction
            try:
                img_path = stripped.split('](')[1].split(')')[0]
                if os.path.exists(img_path):
                    doc.add_picture(img_path, width=Inches(6.0))
                    # caption
                    caption = stripped.split('![')[1].split(']')[0]
                    p = doc.add_paragraph(caption)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.style.font.size = Pt(9)
                    p.style.font.italic = True
            except:
                pass
        elif stripped.startswith('- ') or stripped.startswith('* '):
            add_formatted_para(stripped[2:], 'List Bullet')
        elif stripped.startswith('> '):
            p = add_formatted_para(stripped[2:])
            p.paragraph_format.left_indent = Inches(0.5)
            p.runs[0].italic = True
        else:
            add_formatted_para(stripped)
            
    # output
    out_path = r"c:\Users\G1325\Downloads\Theme Agriculture & Stress Hydrique\Documentation_Agriculture_Precision.docx"
    doc.save(out_path)
    print(f"✅ Docx created at {out_path}")

if __name__ == '__main__':
    create_formatted_doc()
